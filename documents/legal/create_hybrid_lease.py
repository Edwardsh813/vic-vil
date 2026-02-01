#!/usr/bin/env python3
"""
Create a comprehensive HYBRID lease that combines:
- NEW lease format (modern structure, Internet section, updated provisions)
- OLD lease comprehensive rules, fees, and operational details

This creates the definitive Victorian Village lease template.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def copy_paragraph_style(source_para, target_doc):
    """Copy a paragraph with its formatting to the target document"""
    p = target_doc.add_paragraph(source_para.text)

    # Copy paragraph formatting
    if source_para.style:
        try:
            p.style = source_para.style
        except:
            pass

    # Copy runs with formatting
    if len(source_para.runs) > 0:
        p.clear()
        for run in source_para.runs:
            new_run = p.add_run(run.text)
            new_run.bold = run.bold
            new_run.italic = run.italic
            new_run.underline = run.underline
            if run.font.size:
                new_run.font.size = run.font.size
            if run.font.name:
                new_run.font.name = run.font.name

    # Copy alignment
    if source_para.alignment:
        p.alignment = source_para.alignment

    # Copy indentation
    if source_para.paragraph_format.left_indent:
        p.paragraph_format.left_indent = source_para.paragraph_format.left_indent
    if source_para.paragraph_format.first_line_indent:
        p.paragraph_format.first_line_indent = source_para.paragraph_format.first_line_indent

    return p

def create_hybrid_lease():
    """Create the comprehensive hybrid lease"""

    # Load the OLD lease to extract comprehensive sections
    old_doc = Document('victorian-village-lease-template.docx')

    # Create new document
    doc = Document()

    # Set default styles
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(10)

    print("Building hybrid lease document...")

    # ==================== HEADER SECTION ====================
    # Copy header from old lease (paragraphs 0-30 are the header and basic info)
    print("  Adding header...")
    for i in range(0, 31):  # Up to Section 1
        para = old_doc.paragraphs[i]
        if para.text.strip().startswith("1. RENT PAYMENT"):
            break
        copy_paragraph_style(para, doc)

    # ==================== SECTIONS 1-8 ====================
    # These are mostly the same in both, but we'll take from OLD to preserve formatting
    # But we need to update Section 7 to mention Internet separately
    # And ensure Section 8 is the NEW Internet section

    print("  Adding Sections 1-9...")

    # Copy sections 1-7 from OLD lease
    for i in range(31, 38):  # Sections 1-7
        copy_paragraph_style(old_doc.paragraphs[i], doc)

    # Now we need to INSERT the NEW Section 8: INTERNET SERVICE
    # This is a comprehensive section from the NEW lease PDF

    doc.add_paragraph()  # Spacer

    p = doc.add_paragraph()
    run = p.add_run("8. INTERNET SERVICE.")
    run.bold = True

    internet_text = (
        "Internet service is provided to the premises by ERE Fiber LLC (\"ERE Fiber\"). Internet service is REQUIRED for all "
        "units and cannot be opted out of. By occupying the premises, Lessee(s) agrees to the following terms:"
    )
    doc.add_paragraph(internet_text)

    doc.add_paragraph()
    doc.add_paragraph("a. REQUIRED SERVICE AND BILLING: High-speed fiber internet service (500 Mbps) is REQUIRED for all "
                     "units at $45.00 per month. Internet service is billed as a separate line item on the monthly invoice in "
                     "addition to the base rent amount. This is a mandatory utility charge that applies to all units.")

    doc.add_paragraph()
    doc.add_paragraph("b. EQUIPMENT: (i) ERE Fiber provides and owns the Optical Network Terminal (ONT), which is the fiber "
                     "connection device installed at the premises. The ONT remains the property of ERE Fiber and must not be "
                     "tampered with, moved, or removed. (ii) LESSEE IS RESPONSIBLE FOR PROVIDING THEIR OWN WIRELESS ROUTER. "
                     "The ONT provides a wired ethernet connection only. To use WiFi, Lessee(s) must purchase and configure "
                     "their own wireless router. (iii) ERE Fiber is not responsible for Lessee's router, WiFi coverage, or "
                     "connectivity issues caused by Lessee's equipment.")

    doc.add_paragraph()
    doc.add_paragraph("c. SERVICE TERMS: Internet service is subject to ERE Fiber's Terms of Service and Acceptable Use "
                     "Policy, available at erefiber.com. Lessee(s) agrees to comply with all terms and policies.")

    doc.add_paragraph()
    doc.add_paragraph("d. INTERNET SUSPENSION FOR NON-PAYMENT: If rent and/or internet charges remain unpaid after the "
                     "FIFTH (5th) day of the month, Lessor may authorize ERE Fiber to suspend internet service. Lessee(s) will "
                     "receive 72 hours advance written notice before any suspension. Service will be restored within 24 hours of "
                     "payment. Suspension does not waive Lessee's obligation to pay for internet service during suspension.")

    doc.add_paragraph()
    doc.add_paragraph("e. RESIDENTIAL USE ONLY: Internet service is for personal, residential use only. The following activities "
                     "are strictly prohibited: (i) Operating public-facing servers (web, email, game, FTP, or media servers); (ii) "
                     "Running commercial online services or applications; (iii) Reselling internet access to third parties; (iv) "
                     "Operating large-scale file sharing or streaming services; (v) Running cryptocurrency mining operations.")

    doc.add_paragraph()
    doc.add_paragraph("f. ILLEGAL ACTIVITIES PROHIBITED: Lessee(s) shall not use internet service for any illegal purpose, "
                     "including but not limited to: (i) Copyright infringement, including unauthorized downloading or distribution of "
                     "copyrighted materials; (ii) Unauthorized access to computers, networks, or accounts; (iii) Distribution of "
                     "malware, viruses, or malicious software; (iv) Sending spam or phishing emails; (v) Any activity that violates "
                     "federal, state, or local law.")

    doc.add_paragraph()
    doc.add_paragraph("g. SPEED UPGRADES: Lessee(s) may request speed upgrades (1 Gbps or 2 Gbps) by submitting a "
                     "maintenance request. Upgrades are subject to availability and additional monthly charges will be added to "
                     "the invoice.")

    doc.add_paragraph()
    doc.add_paragraph("h. CONSEQUENCES OF VIOLATIONS: Violations of internet service terms may result in: (i) Warning "
                     "requiring immediate corrective action; (ii) Suspension or termination of internet service; (iii) Lessee(s) being "
                     "held financially responsible for investigation costs, damages, and legal fees; (iv) Referral to law "
                     "enforcement for illegal activities; (v) Termination of this Lease at Lessor's discretion.")

    doc.add_paragraph()
    doc.add_paragraph("i. SUPPORT: For internet service issues, Lessee(s) should submit a maintenance request through the "
                     "tenant portal. Issues will be forwarded to ERE Fiber for resolution.")

    doc.add_paragraph()

    # ==================== SECTIONS 9-22 ====================
    # Copy from OLD lease but with renumbering (they become 9-22)
    # OLD section 8 becomes NEW section 9, etc.

    print("  Adding Sections 9-22 (renumbered from old lease)...")

    # Find where section 8 starts in OLD lease
    old_section_8_idx = None
    for i, para in enumerate(old_doc.paragraphs):
        if para.text.strip().startswith("8. RIGHT OF ENTRY"):
            old_section_8_idx = i
            break

    # Copy and renumber sections 8-21 from OLD lease (which become 9-22 in NEW)
    if old_section_8_idx:
        section_num = 9
        for i in range(old_section_8_idx, len(old_doc.paragraphs)):
            text = old_doc.paragraphs[i].text.strip()

            # Stop when we hit section 22 in the old lease
            if text.startswith("22. SECURITY DEPOSIT"):
                # Copy this section but keep it as section 22
                copy_paragraph_style(old_doc.paragraphs[i], doc)
                continue
            elif text.startswith("23. RULES"):
                # We've reached the rules section, we'll add it later
                break

            # Check if this is a numbered section that needs renumbering
            if text and text[0].isdigit() and ". " in text[:5]:
                old_num = text.split(".")[0]
                if old_num.isdigit() and int(old_num) >= 8 and int(old_num) <= 21:
                    # Renumber this section
                    new_text = str(section_num) + "." + text.split(".", 1)[1]
                    p = doc.add_paragraph()
                    run = p.add_run(new_text)
                    run.bold = True
                    section_num += 1
                    continue

            # Copy regular paragraph
            copy_paragraph_style(old_doc.paragraphs[i], doc)

    # ==================== SECTION 23: RULES AND REGULATIONS ====================
    print("  Adding Section 23: Rules and Regulations (from old lease)...")

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("23. RULES AND REGULATIONS.")
    run.bold = True

    # Find and copy the comprehensive rules from OLD lease
    rules_start_idx = None
    rules_end_idx = None

    for i, para in enumerate(old_doc.paragraphs):
        if para.text.strip().startswith("23. RULES AND REGULATIONS"):
            rules_start_idx = i + 1  # Start after the heading
        elif rules_start_idx and para.text.strip().startswith("24. FEES"):
            rules_end_idx = i
            break

    if rules_start_idx and rules_end_idx:
        intro_text = old_doc.paragraphs[rules_start_idx].text.strip()
        if intro_text and not intro_text[0].islower():
            doc.add_paragraph(intro_text)

        for i in range(rules_start_idx, rules_end_idx):
            if old_doc.paragraphs[i].text.strip():
                copy_paragraph_style(old_doc.paragraphs[i], doc)

    # ==================== SECTION 24: FEES ====================
    print("  Adding Section 24: Fees (from old lease)...")

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("24. FEES")
    run.bold = True

    # Find and copy the comprehensive fees from OLD lease
    fees_start_idx = None
    fees_end_idx = None

    for i, para in enumerate(old_doc.paragraphs):
        if para.text.strip().startswith("24. FEES") or para.text.strip() == "24. FEES":
            fees_start_idx = i + 1
        elif fees_start_idx and para.text.strip().startswith("25. ELECTRIC SERVICE"):
            fees_end_idx = i
            break

    if fees_start_idx and fees_end_idx:
        doc.add_paragraph("Lessee(s) agrees to pay the following fees when applicable:")
        doc.add_paragraph()

        for i in range(fees_start_idx, fees_end_idx):
            if old_doc.paragraphs[i].text.strip():
                copy_paragraph_style(old_doc.paragraphs[i], doc)

    # ==================== SECTION 25: ELECTRIC SERVICE ====================
    print("  Adding Section 25: Electric Service (from old lease)...")

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("25. ELECTRIC SERVICE.")
    run.bold = True

    # Find and copy electric service section
    elec_start_idx = None
    elec_end_idx = None

    for i, para in enumerate(old_doc.paragraphs):
        if para.text.strip().startswith("25. ELECTRIC SERVICE"):
            elec_start_idx = i + 1
        elif elec_start_idx and para.text.strip().startswith("26. PARKING"):
            elec_end_idx = i
            break

    if elec_start_idx and elec_end_idx:
        for i in range(elec_start_idx, elec_end_idx):
            if old_doc.paragraphs[i].text.strip():
                copy_paragraph_style(old_doc.paragraphs[i], doc)

    # ==================== SECTIONS 26-30 ====================
    # Copy remaining sections from OLD lease (Parking, Severability, Entire Agreement, Bed Bugs)

    print("  Adding Sections 26-30 (from old lease)...")

    parking_idx = None
    for i, para in enumerate(old_doc.paragraphs):
        if para.text.strip().startswith("26. PARKING"):
            parking_idx = i
            break

    if parking_idx:
        # Copy from parking to the end of numbered sections
        for i in range(parking_idx, len(old_doc.paragraphs)):
            text = old_doc.paragraphs[i].text.strip()

            # Stop at signatures or vehicle info
            if "IN WITNESS WHEREOF" in text or "VEHICLE INFORMATION" in text or "Lessor:" in text:
                break

            copy_paragraph_style(old_doc.paragraphs[i], doc)

    # ==================== SIGNATURE SECTION ====================
    print("  Adding signature section...")

    # Find and copy signature section
    sig_idx = None
    for i, para in enumerate(old_doc.paragraphs):
        if "IN WITNESS WHEREOF" in para.text:
            sig_idx = i
            break

    if sig_idx:
        for i in range(sig_idx, len(old_doc.paragraphs)):
            copy_paragraph_style(old_doc.paragraphs[i], doc)

    return doc

if __name__ == "__main__":
    print("=" * 60)
    print("VICTORIAN VILLAGE LEASE - HYBRID GENERATOR")
    print("=" * 60)
    print()

    doc = create_hybrid_lease()

    # Save the hybrid lease
    output_file = 'victorian-village-lease-template.docx'
    doc.save(output_file)

    print()
    print("=" * 60)
    print(f"✓ Hybrid lease created successfully!")
    print(f"  Output: {output_file}")
    print("=" * 60)
    print()
    print("This lease now includes:")
    print("  ✓ Section 8: Comprehensive Internet Service terms (NEW)")
    print("  ✓ Section 23: Complete Rules and Regulations (OLD)")
    print("  ✓ Section 24: Complete Fee Schedule (OLD)")
    print("  ✓ Section 25: Electric Service Requirement (OLD)")
    print("  ✓ Sections 1-7, 9-22, 26-30: Standard provisions")
    print()
